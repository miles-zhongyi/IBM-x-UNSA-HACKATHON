"""
document_reader.py — extract plain text from uploaded documents.

Supported formats:
  - PDF  (.pdf)  — via PyMuPDF (fitz); handles multilingual: CJK, Arabic,
                   Cyrillic, Devanagari, Latin, etc.
  - DOCX (.docx) — via python-docx; reads paragraphs + table cells.
  - DOC  (.doc)  — tries antiword (subprocess), then python-docx (some .doc
                   files are actually OOXML), then a binary-strip fallback.

Public API
----------
    read_document(path: str | Path) -> str
        Returns extracted Unicode text, or raises DocumentReadError.
"""

from __future__ import annotations

import re
import subprocess
from pathlib import Path
from typing import Union

# ---------------------------------------------------------------------------
# Custom exception
# ---------------------------------------------------------------------------


class DocumentReadError(Exception):
    """Raised when no extraction strategy succeeds."""


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}


def read_document(path: Union[str, Path]) -> str:
    """
    Extract text from a document file.

    Args:
        path: Filesystem path to the document.

    Returns:
        Extracted text as a single Unicode string (paragraphs joined with
        newlines).  May be empty string if the document has no selectable
        text (e.g. scanned image-only PDF).

    Raises:
        DocumentReadError: If the format is unsupported or all strategies fail.
    """
    path = Path(path)
    if not path.exists():
        raise DocumentReadError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentReadError(
            f"Unsupported file type '{suffix}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".doc":
        return _read_doc(path)

    raise DocumentReadError(f"Unhandled suffix: {suffix}")  # unreachable


# ---------------------------------------------------------------------------
# PDF — PyMuPDF
# ---------------------------------------------------------------------------

def _read_pdf(path: Path) -> str:
    """Extract text from a PDF using PyMuPDF (fitz).

    PyMuPDF reads Unicode directly from font character maps, so multilingual
    content (CJK, Arabic, Devanagari, Cyrillic …) is preserved correctly
    without any additional encoding step.
    """
    try:
        import fitz  # pymupdf
    except ImportError as e:
        raise DocumentReadError(
            "PyMuPDF is not installed. Run: pip install pymupdf"
        ) from e

    pages: list[str] = []
    try:
        doc = fitz.open(str(path))
        for page in doc:
            pages.append(str(page.get_text("text")))  # Unicode text
        doc.close()
    except Exception as e:
        raise DocumentReadError(f"PDF extraction failed for {path.name}: {e}") from e

    text = "\n".join(pages).strip()
    if not text:
        # Likely a scanned/image-only PDF — warn but don't crash
        return f"[No selectable text found in {path.name}. " \
               "The document may be a scanned image. " \
               "OCR would be required to extract text.]"
    return text


# ---------------------------------------------------------------------------
# DOCX — python-docx
# ---------------------------------------------------------------------------

def _read_docx(path: Path) -> str:
    """Extract text from a .docx file using python-docx.

    Reads body paragraphs and all table cells to capture structured content.
    Multilingual text is handled transparently because .docx is UTF-8 XML.
    """
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise DocumentReadError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from e

    lines: list[str] = []
    try:
        doc = Document(str(path))

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)

        # Table cells (important for lab result tables etc.)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    lines.append("  |  ".join(row_cells))

    except Exception as e:
        raise DocumentReadError(
            f"DOCX extraction failed for {path.name}: {e}"
        ) from e

    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# DOC — legacy binary Word format
# ---------------------------------------------------------------------------

def _read_doc(path: Path) -> str:
    """Extract text from a legacy .doc file.

    Strategy waterfall:
      1. antiword  — CLI tool; fast and accurate for Word 97-2003 .doc files.
      2. python-docx — some files saved as .doc are actually OOXML (Office 2007+).
      3. Binary strip — last resort; extracts printable Unicode from the raw
         bytes by stripping null bytes and control chars.
    """
    # Strategy 1 — antiword (system package)
    text = _doc_via_antiword(path)
    if text:
        return text

    # Strategy 2 — python-docx (handles mis-named OOXML files)
    try:
        return _read_docx(path)
    except Exception:
        pass

    # Strategy 3 — binary extraction
    text = _doc_via_binary_strip(path)
    if text:
        return text

    raise DocumentReadError(
        f"Could not extract text from {path.name}. "
        "Install 'antiword' for reliable .doc support: "
        "  Linux: sudo apt install antiword | Mac: brew install antiword"
    )


def _doc_via_antiword(path: Path) -> str:
    """Run antiword and capture its output."""
    try:
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True,
            timeout=30,
        )
        if result.returncode == 0:
            return result.stdout.decode("utf-8", errors="replace").strip()
    except (FileNotFoundError, subprocess.TimeoutExpired, Exception):
        pass
    return ""


def _doc_via_binary_strip(path: Path) -> str:
    """Best-effort: read raw bytes, decode as CP1252/UTF-16, strip garbage.

    Word .doc files contain a mix of binary structures and UTF-16LE text
    streams.  This is not perfect but extracts enough readable words for
    the LLM to work with when no other tool is available.
    """
    try:
        raw = path.read_bytes()
        # Try UTF-16 LE (most Word internal text streams)
        try:
            text = raw.decode("utf-16-le", errors="ignore")
        except Exception:
            text = raw.decode("cp1252", errors="ignore")

        # Keep only printable characters; collapse runs of whitespace
        cleaned = re.sub(r"[^\x20-\x7E\u00A0-\uFFFF\n\r\t]", " ", text)
        cleaned = re.sub(r"[ \t]{3,}", " ", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

        # Drop lines that are pure noise (< 4 readable word-chars)
        good_lines = [
            ln for ln in cleaned.splitlines()
            if len(re.findall(r"[A-Za-z\u00C0-\u024F\u4E00-\u9FFF]", ln)) >= 4
        ]
        return "\n".join(good_lines).strip()
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Helper — detect file type from magic bytes (not extension)
# ---------------------------------------------------------------------------

_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"   # ZIP-based (Office Open XML)
_DOC_MAGIC = b"\xd0\xcf\x11\xe0"  # Compound Document (OLE2)


def sniff_format(path: Union[str, Path]) -> str:
    """
    Return the canonical extension ('.pdf', '.docx', '.doc') based on the
    file's magic bytes, ignoring whatever extension the user gave it.

    Returns '' if the format is unrecognised.
    """
    path = Path(path)
    try:
        header = path.read_bytes()[:8]
    except Exception:
        return ""

    if header[:4] == _PDF_MAGIC:
        return ".pdf"
    if header[:4] == _DOCX_MAGIC:
        return ".docx"
    if header[:4] == _DOC_MAGIC:
        return ".doc"
    return ""


def read_document_safe(path: Union[str, Path]) -> tuple[str, str]:
    """
    Like read_document() but:
      - auto-detects format from magic bytes (not the file extension).
      - returns (text, warning) instead of raising.

    Returns:
        text    — extracted text (may be empty on failure).
        warning — empty string on success, or an error/warning message.
    """
    path = Path(path)
    real_ext = sniff_format(path)
    suffix = real_ext or path.suffix.lower()

    # Temporarily rename in memory (create a shadow Path with correct suffix)
    # We don't touch the filesystem — just call the right reader directly.
    try:
        if suffix == ".pdf":
            return _read_pdf(path), ""
        if suffix == ".docx":
            return _read_docx(path), ""
        if suffix == ".doc":
            return _read_doc(path), ""
        return "", f"Unsupported format '{suffix}' for file {path.name}"
    except DocumentReadError as e:
        return "", str(e)
    except Exception as e:
        return "", f"Unexpected error reading {path.name}: {e}"
